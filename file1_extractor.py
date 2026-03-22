"""
Tarang 1.0.0.1 — file1_extractor.py
STATELESS: No local file writes whatsoever.
Returns extracted text as string in result["text"].
Express stores it in MongoDB Document.extractedText.
"""

import re
import logging
from pathlib import Path
from datetime import datetime

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [file1_extractor] %(levelname)s — %(message)s"
)
logger = logging.getLogger("file1_extractor")

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


# ── Text cleaning ─────────────────────────────────────────────────────────────

def sanitize_text(raw: str) -> str:
    replacements = {
        "\u2013": "-", "\u2014": "-", "\u2018": "'", "\u2019": "'",
        "\u201c": '"', "\u201d": '"', "\u2026": "...", "\u00a0": " ",
    }
    for char, replacement in replacements.items():
        raw = raw.replace(char, replacement)
    raw = re.sub(r"[^\x09\x0A\x20-\x7E\u00C0-\u024F]", " ", raw)
    lines = [line.strip() for line in raw.splitlines()]
    cleaned, blank_count = [], 0
    for line in lines:
        if line == "":
            blank_count += 1
            if blank_count <= 2:
                cleaned.append(line)
        else:
            blank_count = 0
            cleaned.append(line)
    return "\n".join(cleaned).strip()


def strip_markdown(text: str) -> str:
    text = re.sub(r"```[\s\S]*?```", "", text)
    text = re.sub(r"`[^`]+`", lambda m: m.group().strip("`"), text)
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", text)
    text = re.sub(r"_{1,3}([^_]+)_{1,3}", r"\1", text)
    text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    text = re.sub(r"^\s*[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)
    return text


# ── Format extractors ─────────────────────────────────────────────────────────

def extract_from_pdf(filepath: Path) -> tuple:
    if pdfplumber is None:
        raise ImportError("pdfplumber not installed. Run: pip install pdfplumber")
    full_text, metadata = [], {"pages": 0, "pages_with_text": 0}
    with pdfplumber.open(filepath) as pdf:
        metadata["pages"] = len(pdf.pages)
        for i, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                full_text.append(f"[Page {i}]\n{page_text.strip()}")
                metadata["pages_with_text"] += 1
            else:
                logger.warning(f"Page {i} yielded no text (possibly scanned).")
    return "\n\n".join(full_text), metadata


def extract_from_docx(filepath: Path) -> tuple:
    if DocxDocument is None:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    doc = DocxDocument(filepath)
    sections = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                table_texts.append(" | ".join(row_data))
    if table_texts:
        sections.append("\n[Tables]\n" + "\n".join(table_texts))
    return "\n\n".join(sections), {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)}


def extract_from_txt(filepath: Path) -> tuple:
    try:
        content = filepath.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        content = filepath.read_text(encoding="latin-1")
        logger.warning("Fell back to latin-1 encoding.")
    return content, {"encoding_used": "utf-8 or latin-1"}


def extract_from_md(filepath: Path) -> tuple:
    raw, metadata = extract_from_txt(filepath)
    return strip_markdown(raw), {**metadata, "markdown_stripped": True}


# ── Main entry point ──────────────────────────────────────────────────────────

def extract(filepath: str, save_output: bool = False) -> dict:
    """
    STATELESS entry point.
    save_output is ignored — no files written to disk.
    Text is returned in result["text"].
    Express/bridge stores it in MongoDB Document.extractedText.

    Returns:
        {
            status, text, word_count, char_count,
            format, metadata, timestamp,
            output_path (always None — kept for bridge compatibility)
        }
    """
    filepath = Path(filepath)
    if not filepath.exists():
        return {"status": "error", "error": f"File not found: {filepath}"}

    ext = filepath.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        return {"status": "error", "error": f"Unsupported format '{ext}'. Supported: {SUPPORTED_EXTENSIONS}"}

    logger.info(f"Starting extraction: {filepath.name} ({ext})")

    try:
        if ext == ".pdf":
            raw_text, fmt_meta = extract_from_pdf(filepath)
        elif ext == ".docx":
            raw_text, fmt_meta = extract_from_docx(filepath)
        elif ext == ".txt":
            raw_text, fmt_meta = extract_from_txt(filepath)
        elif ext == ".md":
            raw_text, fmt_meta = extract_from_md(filepath)
        else:
            return {"status": "error", "error": "Unhandled extension."}

        clean_text = sanitize_text(raw_text)
        if not clean_text:
            return {"status": "error", "error": "Extraction produced empty text. Document may be image-based or empty."}

        word_count = len(clean_text.split())
        char_count = len(clean_text)

        logger.info(f"Extraction complete — {word_count} words, {char_count} chars.")

        return {
            "status":      "success",
            "output_path": None,        # no local file — bridge compatibility only
            "text":        clean_text,  # returned to bridge → Express stores in MongoDB
            "word_count":  word_count,
            "char_count":  char_count,
            "format":      ext.lstrip("."),
            "metadata":    fmt_meta,
            "timestamp":   datetime.utcnow().isoformat() + "Z",
        }

    except Exception as e:
        logger.error(f"Extraction failed: {e}", exc_info=True)
        return {"status": "error", "error": str(e)}


# ── CLI ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python file1_extractor.py <document>")
        sys.exit(1)
    result = extract(sys.argv[1])
    if result["status"] == "success":
        print(f"\n✓ Extraction successful")
        print(f"  Format : {result['format'].upper()}")
        print(f"  Words  : {result['word_count']:,}")
        print(f"  Chars  : {result['char_count']:,}")
        print(f"\nFirst 300 chars:\n{result['text'][:300]}")
    else:
        print(f"\n✗ {result['error']}")
        sys.exit(1)
